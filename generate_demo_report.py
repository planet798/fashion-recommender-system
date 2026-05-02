import json
import os
import sqlite3
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright
from sentence_transformers import SentenceTransformer

from cold_start import ColdStartRecommender
from data_config import config
from interest_updater import InterestUpdater
from user_auth import UserAuth
from user_feedback import UserFeedback


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results" / "demo_assets"
SCREENSHOT_DIR = RESULTS_DIR / "screenshots"
HTML_DIR = RESULTS_DIR / "html"
REPORT_PATH = RESULTS_DIR / "阶段成果_显性个性化演示版.docx"
DATA_PATH = RESULTS_DIR / "demo_metrics.json"
EDGE_PATHS = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
]

DEMO_USERS = [
    {
        "nickname": "demo_fashion",
        "password": "Demo123456",
        "interests": ["Casual", "Formal", "Beauty", "Bohemian"],
        "likes": [("B091TCFBG9", 5), ("B08VNGMWRQ", 4), ("B09P4XGS2W", 4)],
        "dislikes": [("B07JWLTCLX", 1)],
        "query": "生日礼物",
        "demo_order": 1,
        "profile_label": "时尚礼物型用户",
    },
    {
        "nickname": "demo_tech",
        "password": "Demo123456",
        "interests": ["Gadgets", "Audio", "Wearables", "Gaming"],
        "likes": [("B07JWLTCLX", 5), ("B00AAFL08G", 4), ("B07DWD57H2", 4)],
        "dislikes": [("B091TCFBG9", 1)],
        "query": "生日礼物",
        "demo_order": 2,
        "profile_label": "数码音频型用户",
    },
    {
        "nickname": "demo_home",
        "password": "Demo123456",
        "interests": ["Home Decor", "Kitchen", "Travel"],
        "likes": [("B081K8RVSR", 5), ("B08KP6YM3T", 4), ("B016RO0F3O", 4)],
        "dislikes": [("B07DWD57H2", 1)],
        "query": "生日礼物",
        "demo_order": 3,
        "profile_label": "家居装饰型用户",
    },
]


def ensure_dirs():
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    HTML_DIR.mkdir(parents=True, exist_ok=True)


def get_edge_executable():
    for path in EDGE_PATHS:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("Edge or Chrome executable not found.")


def ensure_demo_users():
    auth = UserAuth()
    feedback = UserFeedback()
    updater = InterestUpdater()
    db_path = ROOT / "data" / "user_recommendation.db"
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()

    rows = []
    for cfg in DEMO_USERS:
        success, _, user_id = auth.login(cfg["nickname"], cfg["password"])
        if not success:
            success, msg, user_id = auth.register(cfg["nickname"], cfg["password"], cfg["interests"])
            if not success:
                raise RuntimeError(f"Failed to create demo user {cfg['nickname']}: {msg}")

        auth.update_interests(user_id, cfg["interests"])
        cur.execute("DELETE FROM user_feedback WHERE user_id = ?", (user_id,))
        cur.execute("DELETE FROM exposure_history WHERE user_id = ?", (user_id,))
        conn.commit()

        for item_id, rating in cfg["likes"]:
            feedback.add_feedback(user_id, item_id, liked=True, rating=rating)
            updater.update_interest_weights(user_id, [{"item_id": item_id, "liked": True, "rating": rating}])
            feedback.record_exposure(user_id, [item_id])

        for item_id, rating in cfg["dislikes"]:
            feedback.add_feedback(user_id, item_id, liked=False, rating=rating)
            updater.update_interest_weights(user_id, [{"item_id": item_id, "liked": False, "rating": rating}])
            feedback.record_exposure(user_id, [item_id])

        cur.execute(
            "SELECT interest_tag, weight FROM user_interests WHERE user_id = ? ORDER BY weight DESC LIMIT 5",
            (user_id,),
        )
        top_interests = cur.fetchall()
        rows.append({
            **cfg,
            "user_id": user_id,
            "top_interests": top_interests,
            "like_count": len(cfg["likes"]),
            "dislike_count": len(cfg["dislikes"]),
        })

    conn.close()
    return rows


def load_materials():
    items = pd.read_csv(config.items_csv)
    items["item_id"] = items["item_id"].astype(str)
    text_features = {
        str(item_id): np.array(vec, dtype=np.float32)
        for item_id, vec in np.load(config.text_features, allow_pickle=True).item().items()
    }
    model = SentenceTransformer("models/paraphrase-MiniLM-L3-v2")
    conn = sqlite3.connect(ROOT / "data" / "user_recommendation.db")
    cur = conn.cursor()
    cur.execute("SELECT category, tag FROM interest_tags ORDER BY category, tag")
    interest_tags = {}
    for category, tag in cur.fetchall():
        interest_tags.setdefault(category, []).append(tag)
    conn.close()
    recommender = ColdStartRecommender(items_df=items, text_features=text_features, interest_tags=interest_tags)
    return items, text_features, model, recommender


def get_interest_weights(user_id):
    conn = sqlite3.connect(ROOT / "data" / "user_recommendation.db")
    cur = conn.cursor()
    cur.execute(
        "SELECT interest_tag, weight FROM user_interests WHERE user_id = ? ORDER BY weight DESC",
        (user_id,),
    )
    rows = {tag: float(weight) for tag, weight in cur.fetchall()}
    conn.close()
    return rows


def get_negative_interest_weights(user_id, recommender):
    conn = sqlite3.connect(ROOT / "data" / "user_recommendation.db")
    cur = conn.cursor()
    cur.execute("SELECT item_id, liked, rating FROM user_feedback WHERE user_id = ?", (user_id,))
    negatives = {}
    for item_id, liked, rating in cur.fetchall():
        is_negative = liked == 0 or liked is False or (rating is not None and rating <= 2)
        if not is_negative:
            continue
        for tag, score in recommender.get_item_interest_matches(str(item_id), topn=3, min_score=0.0):
            negatives[tag] = negatives.get(tag, 0.0) + float(score)
    conn.close()
    if not negatives:
        return {}
    max_score = max(negatives.values()) or 1.0
    return {tag: score / max_score for tag, score in negatives.items()}


def compute_experiment_metrics(user_rows):
    items, text_features, model, recommender = load_materials()
    item_ids = np.array(list(text_features.keys()))
    vectors = np.array(
        [text_features[item_id] / (np.linalg.norm(text_features[item_id]) + 1e-8) for item_id in item_ids],
        dtype=np.float32,
    )
    title_map = dict(zip(items["item_id"], items["title"].fillna("")))

    metrics = {}
    for row in user_rows:
        query = row["query"]
        query_vec = model.encode(query, normalize_embeddings=True)
        base_scores = np.dot(vectors, query_vec)
        candidate_idx = np.argsort(base_scores)[::-1][:50]
        candidate_ids = item_ids[candidate_idx]

        generic_top = []
        for idx in candidate_idx[:8]:
            item_id = str(item_ids[idx])
            generic_top.append({
                "item_id": item_id,
                "title": title_map.get(item_id, item_id),
                "generic_score": float(base_scores[idx]),
            })

        interest_weights = get_interest_weights(row["user_id"])
        negative_weights = get_negative_interest_weights(row["user_id"], recommender)
        interest_raw = {
            str(item_id): recommender.score_item_against_interests(str(item_id), interest_weights)
            for item_id in candidate_ids
        }
        negative_raw = {
            str(item_id): recommender.score_item_against_interests(str(item_id), negative_weights)
            for item_id in candidate_ids
        }
        max_interest = max(interest_raw.values()) or 1.0
        max_negative = max(negative_raw.values()) or 1.0

        personalized = []
        for idx in candidate_idx:
            item_id = str(item_ids[idx])
            base_score = float(base_scores[idx])
            interest_score = float(interest_raw.get(item_id, 0.0)) / max_interest if max_interest > 0 else 0.0
            negative_score = float(negative_raw.get(item_id, 0.0)) / max_negative if max_negative > 0 else 0.0
            final_score = 0.64 * base_score + 0.31 * interest_score - 0.17 * negative_score
            personalized.append({
                "item_id": item_id,
                "title": title_map.get(item_id, item_id),
                "generic_score": base_score,
                "personalization_score": interest_score,
                "negative_score": negative_score,
                "final_score": final_score,
            })
        personalized.sort(key=lambda x: x["final_score"], reverse=True)
        personalized_top = personalized[:8]

        generic_rank = {entry["item_id"]: idx + 1 for idx, entry in enumerate(generic_top)}
        personalized_rank = {entry["item_id"]: idx + 1 for idx, entry in enumerate(personalized_top)}
        changed_items = sum(1 for item_id, rank in personalized_rank.items() if generic_rank.get(item_id) != rank)
        overlap_count = len(set(generic_rank.keys()) & set(personalized_rank.keys()))

        metrics[row["nickname"]] = {
            "query": query,
            "changed_items": changed_items,
            "top1_changed": bool(generic_top and personalized_top and generic_top[0]["item_id"] != personalized_top[0]["item_id"]),
            "overlap_count": overlap_count,
            "generic_top": generic_top,
            "personalized_top": personalized_top,
        }
    return metrics


def render_login_html():
    html = """<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>login</title>
<style>
body{margin:0;font-family:'Segoe UI','Microsoft YaHei',sans-serif;background:linear-gradient(180deg,#f7fbff,#eaf2fc);color:#102544}
.shell{display:grid;grid-template-columns:1.2fr 0.8fr;gap:22px;padding:32px;min-height:100vh;box-sizing:border-box}
.hero{border-radius:28px;padding:28px;background:linear-gradient(160deg,rgba(8,20,44,.96),rgba(12,41,84,.9));color:#f2f7ff;position:relative;overflow:hidden}
.hero:before{content:'';position:absolute;inset:0;background-image:linear-gradient(rgba(129,194,255,.08) 1px,transparent 1px),linear-gradient(90deg,rgba(129,194,255,.08) 1px,transparent 1px);background-size:34px 34px}
.eyebrow{display:inline-block;padding:6px 12px;border-radius:999px;background:rgba(74,154,255,.12);font-size:12px;font-weight:700;letter-spacing:.08em;text-transform:uppercase}
.title{font-size:46px;font-weight:800;line-height:1.05;max-width:11ch;margin:18px 0}
.text{font-size:16px;line-height:1.8;max-width:760px;color:rgba(238,245,255,.85)}
.grid{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-top:26px}
.card{border:1px solid rgba(120,180,255,.18);background:rgba(255,255,255,.08);padding:16px;border-radius:20px}
.card strong{display:block;font-size:18px;margin-bottom:6px}
.panel{background:rgba(255,255,255,.9);border:1px solid rgba(104,149,219,.16);border-radius:28px;padding:28px;box-shadow:0 22px 50px rgba(31,73,135,.12)}
.input{background:#fff;border:1px solid rgba(89,132,199,.18);border-radius:14px;padding:14px 16px;margin:10px 0 16px}
.btn{margin-top:8px;background:linear-gradient(135deg,#2d7ff9,#1557c8);color:#fff;border:none;border-radius:14px;padding:14px 16px;font-weight:700;text-align:center}
.signal{margin-top:20px;border-radius:22px;background:linear-gradient(160deg,#07162d,#0d2f5c);height:220px;color:#eff7ff;padding:18px 20px;box-sizing:border-box}
.signal .big{font-size:28px;font-weight:800;margin-top:20px}
</style></head><body>
<div class="shell">
<div class="hero">
<span class="eyebrow">Personalization Gateway</span>
<div class="title">先建立用户画像，再进入推荐工作台。</div>
<div class="text">系统会在登录后绑定兴趣标签、点赞/点踩与评分信号，再把这些信号注入多模态召回和排序流程。这样进入搜索界面后，评委可以直接看到同一查询在不同用户下的结果差异。</div>
<div class="grid">
<div class="card"><strong>Login First</strong>搜索入口仅在认证后开放，避免账号与搜索混杂。</div>
<div class="card"><strong>Visible Persona</strong>兴趣标签、正负反馈和重排提升会直接展示。</div>
<div class="card"><strong>Demo Ready</strong>支持通用排序与个性化排序对照演示。</div>
</div>
</div>
<div>
<div class="panel">
<span class="eyebrow" style="color:#0b58c7;background:rgba(74,154,255,.12)">Access Control</span>
<h2>登录 / 注册</h2>
<div class="input">昵称：例如 demo_fashion</div>
<div class="input">密码：••••••••••••</div>
<div class="btn">进入推荐工作台</div>
</div>
<div class="signal"><div style="font-size:12px;letter-spacing:.16em;text-transform:uppercase;color:#7fc3ff;">Live Signal Console</div><div class="big">Visible Personalization Online</div></div>
</div>
</div></body></html>"""
    path = HTML_DIR / "login_mock.html"
    path.write_text(html, encoding="utf-8")
    return path


def render_profile_html(user_row):
    cards = "".join(
        f"<div class='profile'><strong>{tag}</strong><div>权重 {weight:.2f}</div></div>"
        for tag, weight in user_row["top_interests"]
    )
    html = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>profile</title>
<style>
body{{margin:0;font-family:'Segoe UI','Microsoft YaHei',sans-serif;background:linear-gradient(180deg,#f7fbff,#eef5ff);color:#102544}}
.wrap{{padding:28px}}
.banner{{background:linear-gradient(135deg,rgba(11,36,70,.96),rgba(10,69,133,.92));color:#eef6ff;border-radius:24px;padding:22px}}
.panel{{margin-top:20px;background:rgba(255,255,255,.92);border:1px solid rgba(85,132,205,.14);border-radius:24px;padding:24px}}
.grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-top:16px}}
.profile{{border:1px solid rgba(111,166,235,.14);border-radius:20px;padding:18px;background:#fff}}
.profile strong{{display:block;font-size:20px;margin-bottom:8px}}
</style></head><body><div class="wrap">
<div class="banner"><div style="font-size:12px;letter-spacing:.14em;text-transform:uppercase">Profile Snapshot</div><h1>当前用户画像</h1><p>账号：{user_row['nickname']} | 定位：{user_row['profile_label']}</p></div>
<div class="panel"><h2>高权重兴趣标签</h2><div class="grid">{cards}</div></div>
</div></body></html>"""
    path = HTML_DIR / f"{user_row['nickname']}_profile.html"
    path.write_text(html, encoding="utf-8")
    return path


def render_lab_html(user_row, metric):
    def card_html(entry, personalized=False):
        meta = f"Generic {entry['generic_score']:.3f}"
        if personalized:
            meta += f" | Personalization {entry['personalization_score']:.3f} | Negative {entry['negative_score']:.3f}"
        return f"<div class='card'><strong>{entry['title'][:110]}</strong><div class='sub'>Item ID: {entry['item_id']}</div><div class='meta'>{meta}</div></div>"

    generic_cards = "".join(card_html(entry, False) for entry in metric["generic_top"][:4])
    personalized_cards = "".join(card_html(entry, True) for entry in metric["personalized_top"][:4])
    html = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>lab</title>
<style>
body{{margin:0;font-family:'Segoe UI','Microsoft YaHei',sans-serif;background:linear-gradient(180deg,#f7fbff,#eef5ff);color:#102544}}
.wrap{{padding:24px}}
.banner{{background:linear-gradient(135deg,rgba(11,36,70,.96),rgba(10,69,133,.92));color:#eef6ff;border-radius:24px;padding:20px 24px}}
.eyebrow{{display:inline-block;padding:6px 12px;border-radius:999px;background:rgba(255,255,255,.1);font-size:12px;letter-spacing:.08em;text-transform:uppercase}}
.lab{{margin-top:18px;background:#fff;border-radius:24px;padding:22px;box-shadow:0 18px 44px rgba(51,102,173,.10)}}
.metrics{{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin:16px 0}}
.metric{{background:rgba(45,127,249,.08);border:1px solid rgba(87,132,198,.12);border-radius:18px;padding:16px}}
.metric .label{{font-size:13px;color:#4f678a}}
.metric .value{{font-size:30px;font-weight:800;margin-top:8px}}
.cols{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}
.section{{background:rgba(255,255,255,.88);border:1px solid rgba(87,132,198,.12);border-radius:22px;padding:18px}}
.card{{background:rgba(255,255,255,.94);border:1px solid rgba(87,132,198,.12);border-radius:18px;padding:14px;margin-bottom:12px}}
.card strong{{display:block;font-size:18px;line-height:1.45}}
.sub{{font-size:12px;color:#5f7391;margin-top:6px}}
.meta{{font-size:13px;color:#1f67de;margin-top:8px}}
</style></head><body><div class="wrap">
<div class="banner"><span class="eyebrow">Visible Personalization Lab</span><h1>{user_row['nickname']} | {user_row['profile_label']}</h1><p>固定查询：{metric['query']}。当前页面用于展示同一查询在不同画像下的推荐差异。</p></div>
<div class="lab">
<h2>显性个性化实验台</h2>
<div class="metrics">
<div class="metric"><div class="label">排序变化商品数</div><div class="value">{metric['changed_items']}</div></div>
<div class="metric"><div class="label">Top1 是否变化</div><div class="value">{'是' if metric['top1_changed'] else '否'}</div></div>
<div class="metric"><div class="label">结果重合数</div><div class="value">{metric['overlap_count']}</div></div>
</div>
<div class="cols">
<div class="section"><h3>通用排序</h3>{generic_cards}</div>
<div class="section"><h3>个性化排序</h3>{personalized_cards}</div>
</div>
</div></div></body></html>"""
    path = HTML_DIR / f"{user_row['nickname']}_lab.html"
    path.write_text(html, encoding="utf-8")
    return path


def capture_static_assets(user_rows, experiment_metrics):
    assets = {"screenshots": {}, "lab_metrics": experiment_metrics}
    edge_path = get_edge_executable()

    with sync_playwright() as p:
        browser = p.chromium.launch(executable_path=edge_path, headless=True)
        context = browser.new_context(viewport={"width": 1600, "height": 1800}, locale="zh-CN")

        login_html = render_login_html()
        page = context.new_page()
        page.goto(login_html.as_uri(), wait_until="load", timeout=120000)
        login_path = SCREENSHOT_DIR / "login_page.png"
        page.screenshot(path=str(login_path), full_page=True)
        assets["screenshots"]["login_page"] = str(login_path)
        page.close()

        profile_html = render_profile_html(user_rows[0])
        page = context.new_page()
        page.goto(profile_html.as_uri(), wait_until="load", timeout=120000)
        profile_path = SCREENSHOT_DIR / f"{user_rows[0]['nickname']}_profile.png"
        page.screenshot(path=str(profile_path), full_page=True)
        assets["screenshots"]["profile_page"] = str(profile_path)
        page.close()

        for row in user_rows:
            lab_html = render_lab_html(row, experiment_metrics[row["nickname"]])
            page = context.new_page()
            page.goto(lab_html.as_uri(), wait_until="load", timeout=120000)
            shot_path = SCREENSHOT_DIR / f"{row['nickname']}_lab.png"
            page.screenshot(path=str(shot_path), full_page=True)
            assets["screenshots"][row["nickname"]] = str(shot_path)
            page.close()

        context.close()
        browser.close()

    return assets


def summarize_lab_metrics(experiment_metrics):
    rows = []
    for cfg in DEMO_USERS:
        metric = experiment_metrics[cfg["nickname"]]
        rows.append({
            "nickname": cfg["nickname"],
            "query": cfg["query"],
            "排序变化商品数": metric["changed_items"],
            "Top1 是否变化": "是" if metric["top1_changed"] else "否",
            "结果重合数": metric["overlap_count"],
        })
    return rows


def build_doc(user_rows, assets):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    run = title.add_run("阶段成果报告：显性个性化演示版")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph(
        "项目：多模态个性化推荐系统\n"
        "日期：2026-04-30\n"
        "当前阶段目标：完成登录-工作台分离、显性个性化展示链路、答辩演示用户与阶段成果材料固化。"
    )

    doc.add_heading("1. 本阶段完成内容", level=1)
    for text in [
        "界面调整为先登录后进入工作台，登录区与搜索区彻底分离。",
        "新增科技风视觉层：渐变背景、玻璃拟态卡片、动态信号区、统一按钮与指标标签。",
        "新增显性个性化实验台，可同时展示通用排序与个性化排序。",
        "补齐负向反馈建模，点踩与低评分会参与兴趣抑制。",
        "固化 3 个答辩测试账号、截图素材与阶段成果文档生成脚本。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("2. 答辩测试用户方案", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "顺序"
    hdr[1].text = "账号"
    hdr[2].text = "密码"
    hdr[3].text = "画像定位"
    hdr[4].text = "核心兴趣"
    hdr[5].text = "建议操作"
    for row in user_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["demo_order"])
        cells[1].text = row["nickname"]
        cells[2].text = row["password"]
        cells[3].text = row["profile_label"]
        cells[4].text = ", ".join(row["interests"])
        ops = [f"喜欢 {item_id}" for item_id, _ in row["likes"][:2]] + [f"点踩 {row['dislikes'][0][0]}"]
        cells[5].text = "；".join(ops)
    doc.add_paragraph("推荐答辩演示顺序：demo_fashion -> demo_tech -> demo_home，三次都使用同一查询“生日礼物”。")

    doc.add_heading("3. 实验数据摘录", level=1)
    metric_table = doc.add_table(rows=1, cols=5)
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = metric_table.rows[0].cells
    hdr[0].text = "账号"
    hdr[1].text = "查询"
    hdr[2].text = "排序变化商品数"
    hdr[3].text = "Top1 是否变化"
    hdr[4].text = "结果重合数"
    for row in summarize_lab_metrics(assets["lab_metrics"]):
        cells = metric_table.add_row().cells
        cells[0].text = row["nickname"]
        cells[1].text = row["query"]
        cells[2].text = str(row["排序变化商品数"])
        cells[3].text = row["Top1 是否变化"]
        cells[4].text = str(row["结果重合数"])
    doc.add_paragraph("说明：排序变化商品数越高，表示个性化对排序的可见干预越明显；Top1 变化用于答辩时直观说明“同问不同答”。")

    doc.add_heading("4. 当前用户画像权重", level=1)
    profile_table = doc.add_table(rows=1, cols=3)
    hdr = profile_table.rows[0].cells
    hdr[0].text = "账号"
    hdr[1].text = "前5兴趣"
    hdr[2].text = "反馈概况"
    for row in user_rows:
        cells = profile_table.add_row().cells
        cells[0].text = row["nickname"]
        cells[1].text = "；".join(f"{tag}({weight:.2f})" for tag, weight in row["top_interests"])
        cells[2].text = f"喜欢 {row['like_count']} 项，负反馈 {row['dislike_count']} 项"

    doc.add_heading("5. 演示界面截图", level=1)
    for key, caption in [
        ("login_page", "图1 登录页：先登录后进入推荐工作台"),
        ("profile_page", "图2 用户画像页：展示高权重兴趣"),
        ("demo_fashion", "图3 demo_fashion 的显性个性化实验台"),
        ("demo_tech", "图4 demo_tech 的显性个性化实验台"),
        ("demo_home", "图5 demo_home 的显性个性化实验台"),
    ]:
        if key not in assets["screenshots"]:
            continue
        doc.add_paragraph(caption)
        doc.add_picture(assets["screenshots"][key], width=Inches(6.5))

    doc.add_heading("6. 答辩建议话术", level=1)
    for text in [
        "先展示登录页，说明系统已改成先认证后进入工作台，用户身份会持续影响推荐。",
        "登录 demo_fashion，用查询“生日礼物”展示偏时尚礼物方向的排序。",
        "切换到 demo_tech，保持同一查询，展示数码音频方向的推荐结果变化。",
        "切换到 demo_home，再次保持同一查询，展示家居装饰方向的变化。",
        "最后现场执行一次点踩或低评分，再刷新实验台，展示负向偏好抑制后的重排效果。",
    ]:
        doc.add_paragraph(text, style="List Number")

    doc.save(REPORT_PATH)


def main():
    ensure_dirs()
    user_rows = ensure_demo_users()
    experiment_metrics = compute_experiment_metrics(user_rows)
    assets = capture_static_assets(user_rows, experiment_metrics)
    output = {"users": user_rows, "lab_metrics": experiment_metrics, "screenshots": assets["screenshots"]}
    with open(DATA_PATH, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    build_doc(user_rows, assets)
    print(f"Report generated: {REPORT_PATH}")
    print(f"Metrics saved: {DATA_PATH}")


if __name__ == "__main__":
    main()
